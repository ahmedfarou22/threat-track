from .update_internal_links import update_internal_links
from .html_2_docx_custom import html_2_sub_docx
from docxtpl import DocxTemplate, InlineImage
from django.template import Template, Context
from assessments.models import Assessment
from datetime import timedelta
from django.forms import model_to_dict
from django.http import HttpResponse
from matplotlib.patches import Patch
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from django.db import models
from docx.shared import Mm
from docx import Document
import jinja2, io, re


# Conver and fix model -> dict
def model_to_dict_recursive(instance):
    data = model_to_dict(instance)
    for field_name, field_value in data.items():
        if isinstance(field_value, models.Model):
            data[field_name] = model_to_dict_recursive(field_value)
        elif isinstance(field_value, list):
            new_list = []
            for item in field_value:
                if isinstance(item, models.Model):
                    new_list.append(model_to_dict_recursive(item))
                else:
                    new_list.append(item)
            data[field_name] = new_list
    return data


def fix_genrated_dict(dict):
    assessment = Assessment.objects.get(id=dict["id"])

    # Remove Assessment ID, Files
    del dict["id"]
    del dict["af_name"]
    del dict["v_fields"]
    del dict["tasks"]
    del dict["priority"]

    # Fix assigned users
    assigned_users = []
    for field in dict["assigned_users"]:
        assigned_users.append(f"{field['first_name']} {field['last_name']} ")
    dict["assigned_users"] = assigned_users

    # Add man days (exculuding weekend)
    total_days = (dict["end_date"] - dict["start_date"]).days
    excluded_days = 0
    current_date = dict["start_date"]
    while current_date <= dict["end_date"]:
        if current_date.weekday() in [4, 5]:
            excluded_days += 1
        current_date += timedelta(days=1)
    man_days = total_days - excluded_days
    man_days += 1
    dict["man_days"] = str(man_days)

    # Fix Dates
    dict["start_date"] = dict["start_date"].strftime("%B %d, %Y")
    dict["end_date"] = dict["end_date"].strftime("%B %d, %Y")

    # Fix who_created, client, status
    dict["who_created"] = (
        assessment.who_created.first_name + " " + assessment.who_created.last_name
    )
    dict["status"] = assessment.status.name
    dict["client"] = model_to_dict(assessment.client) if dict["client"] else None

    # Remove spaces from aditonal fileds keys
    s_modified_dict = {}
    if dict["s_fields"]:
        if dict["s_fields"].items():
            for key, value in dict["s_fields"].items():
                modified_key = key.replace(" ", "_")
                s_modified_dict[modified_key] = value
        dict["s_fields"] = s_modified_dict

    a_modified_dict = {}
    if dict["a_fields"]:
        if dict["a_fields"].items():
            for key, value in dict["a_fields"].items():
                modified_key = key.replace(" ", "_")
                a_modified_dict[modified_key] = value
        dict["a_fields"] = a_modified_dict

    for vuln in dict["vulnerabilities"]:
        if vuln["fields"]:
            if vuln["fields"].items():
                v_modified_dict2 = {}
                for key, value in vuln["fields"].items():
                    modified_key = key.replace(" ", "_")
                    v_modified_dict2[modified_key] = value
                vuln["fields"] = v_modified_dict2

    return dict


def all_in_one(instance):
    return fix_genrated_dict(model_to_dict_recursive(instance))


# Custom Jinja functions
def sort_cvss(vulnerabilities, high_or_low=None):
    if high_or_low == "high":
        vulnerabilities = sorted(vulnerabilities, key=lambda x: x["cvss"], reverse=True)
    elif high_or_low == "low":
        vulnerabilities = sorted(vulnerabilities, key=lambda x: x["cvss"])
    else:
        vulnerabilities = sorted(vulnerabilities, key=lambda x: x["cvss"], reverse=True)

    # Reorder 'number' key starting from 1
    for i, vulnerability in enumerate(vulnerabilities):
        vulnerability["number"] = i + 1

    return vulnerabilities


def count_risk_rating(vulnerabilities, what):
    count = [v for v in vulnerabilities if v["risk_rating"] == what.capitalize()]
    return len(count)


def count_all(vulnerabilities):
    return len(vulnerabilities)


def genrate_report(template_obj, context):
    try:
        template_path = template_obj.file
        chart_settings = template_obj.chart_settings

        # Check if chart_settings is None or doesn't have 'charts' key
        if chart_settings is None:
            chart_settings = {"charts": []}
        elif not isinstance(chart_settings, dict):
            chart_settings = {"charts": []}
        elif "charts" not in chart_settings:
            chart_settings["charts"] = []
        elif chart_settings["charts"] is None:
            chart_settings["charts"] = []

        # Set defaults for potentially missing context fields BEFORE any access
        if context.get("a_fields") is None:
            context["a_fields"] = {}
        if context.get("s_fields") is None:
            context["s_fields"] = {}
        if context.get("vulnerabilities") is None:
            context["vulnerabilities"] = []
        if context.get("files") is None:
            context["files"] = []
        if context.get("assigned_users") is None:
            context["assigned_users"] = []

        # Prevent Template Injection. (Works by getting all the strings inside the word document. runs secure django jinja rendering if it passes it is secure)
        if True:
            strings_all = ""
            test_doc = Document(template_path)
            for paragraph in test_doc.paragraphs:
                strings_all = strings_all + str(paragraph.text)
            for table in test_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            strings_all = strings_all + str(paragraph.text)

            for pattern in [
                r"\|sort_cvss\(.*?\)",
                r"\|count_risk_rating\(.*?\)",
                r"\|count_all\(.*?\)",
                r"{% cellbg .*? %}",
            ]:
                strings_all = re.sub(pattern, "", strings_all)
            strings_all = (
                strings_all.replace("{%p", "{%")
                .replace("{%tr", "{%")
                .replace("{%tc", "{%")
                .replace("{%r", "{%")
                .replace("{{p", "{{")
                .replace("{{r", "{{")
            )

            strings_to_test = Template(str(strings_all))
            context_tmp = Context(context)
            strings_to_test.render(context_tmp)

        doc = DocxTemplate(template_path)

        # Custom Jinja filters
        jinja_env = jinja2.Environment()
        jinja_env.filters["sort_cvss"] = sort_cvss
        jinja_env.filters["count_risk_rating"] = count_risk_rating
        jinja_env.filters["count_all"] = count_all

        # Render Charts
        for chart_setting in chart_settings["charts"]:
            chart_type = chart_setting["type"]
            labels = []
            colors = []
            sizes = []
            for level in ["Critical", "High", "Medium", "Low"]:
                count = count_risk_rating(context["vulnerabilities"], level)
                if count > 0:
                    labels.append(level)
                    colors.append(chart_setting["colors"][level])
                    sizes.append(count)

            # Generate chart
            fig, ax = plt.subplots()
            if chart_type == "pie":
                ax.pie(
                    sizes,
                    labels=labels,
                    colors=colors,
                    autopct="%1.1f%%",
                    startangle=90,
                )
                legend_elements = [
                    Patch(facecolor=color, edgecolor="black", label=label)
                    for label, color in zip(labels, colors)
                ]
                ax.legend(handles=legend_elements, loc="upper left")
            elif chart_type == "bar":
                ax.bar(labels, sizes, color=colors)
                ax.set_xlabel(chart_setting["x_label"])
                ax.set_ylabel(chart_setting["y_label"])

            ax.set_title(chart_setting["title"])

            try:
                buffer = io.BytesIO()
                fig.savefig(buffer, format="png")
                plt.close(fig)
                image_bytes = buffer.getvalue()

                context[str(chart_type) + "_chart_bytes"] = image_bytes
                context[str(chart_type) + "_chart_size"] = chart_setting.get(
                    "size", 120
                )
            except Exception as e:
                plt.close(fig)  # Ensure figure is closed even on error
                context[str(chart_type) + "_chart"] = None

        # Rendering Additional Fields
        if context["s_fields"]:
            for filed in context["s_fields"]:
                try:
                    text_or_subdoc = html_2_sub_docx(context["s_fields"][filed])
                    context["s_fields"][filed] = (
                        text_or_subdoc
                        if isinstance(text_or_subdoc, str)
                        else doc.new_subdoc(text_or_subdoc)
                    )
                except Exception as e:
                    print(f"Error rendering summary field {filed}: {str(e)}")
                    context["s_fields"][filed] = str(
                        context["s_fields"][filed]
                    )  # Fallback to plain text

        if context["a_fields"]:
            for filed in context["a_fields"]:
                try:
                    text_or_subdoc = html_2_sub_docx(context["a_fields"][filed])
                    context["a_fields"][filed] = (
                        text_or_subdoc
                        if isinstance(text_or_subdoc, str)
                        else doc.new_subdoc(text_or_subdoc)
                    )
                except Exception as e:
                    print(f"Error rendering additional field {filed}: {str(e)}")
                    context["a_fields"][filed] = str(
                        context["a_fields"][filed]
                    )  # Fallback to plain text

        # Rendering vulnerabilities
        for vuln in context["vulnerabilities"]:
            try:
                text_or_subdoc = html_2_sub_docx(vuln["description"])
                vuln["description"] = (
                    text_or_subdoc
                    if isinstance(text_or_subdoc, str)
                    else doc.new_subdoc(text_or_subdoc)
                )
            except Exception as e:
                print(f"Error rendering vulnerability description: {str(e)}")
                vuln["description"] = str(vuln["description"])

            try:
                text_or_subdoc = html_2_sub_docx(vuln["impact"])
                vuln["impact"] = (
                    text_or_subdoc
                    if isinstance(text_or_subdoc, str)
                    else doc.new_subdoc(text_or_subdoc)
                )
            except Exception as e:
                print(f"Error rendering vulnerability impact: {str(e)}")
                vuln["impact"] = str(vuln["impact"])

            try:
                text_or_subdoc = html_2_sub_docx(vuln["remediation"])
                vuln["remediation"] = (
                    text_or_subdoc
                    if isinstance(text_or_subdoc, str)
                    else doc.new_subdoc(text_or_subdoc)
                )
            except Exception as e:
                print(f"Error rendering vulnerability remediation: {str(e)}")
                vuln["remediation"] = str(vuln["remediation"])

            try:
                text_or_subdoc = html_2_sub_docx(vuln["poc_text"])
                vuln["poc_text"] = (
                    text_or_subdoc
                    if isinstance(text_or_subdoc, str)
                    else doc.new_subdoc(text_or_subdoc)
                )
            except Exception as e:
                print(f"Error rendering vulnerability poc_text: {str(e)}")
                vuln["poc_text"] = str(vuln["poc_text"])

        for vuln in context["vulnerabilities"]:
            if vuln["fields"]:
                for field in vuln["fields"]:
                    try:
                        text_or_subdoc = html_2_sub_docx(vuln["fields"][field])
                        vuln["fields"][field] = (
                            text_or_subdoc
                            if isinstance(text_or_subdoc, str)
                            else doc.new_subdoc(text_or_subdoc)
                        )
                    except Exception as e:
                        print(f"Error rendering vulnerability field {field}: {str(e)}")
                        vuln["fields"][field] = str(
                            vuln["fields"][field]
                        )  # Fallback to plain text

        # Rendering client logo
        try:
            # Check if client exists and has a valid logo file
            if (
                context.get("client") is not None
                and isinstance(context["client"], dict)
                and context["client"].get("logo") is not None
            ):
                logo_field = context["client"]["logo"]

                # Check if ImageFieldFile has an actual file
                if (
                    hasattr(logo_field, "name")
                    and logo_field.name
                    and str(logo_field.name).strip()
                ):
                    context["client"]["logo"] = InlineImage(
                        doc, context["client"]["logo"], width=Mm(20)
                    )
                else:
                    context["client"]["logo"] = None
            else:
                # Ensure client dict exists before setting logo to None
                if context.get("client") is None:
                    context["client"] = {}
                context["client"]["logo"] = None
        except Exception as e:
            # Ensure client dict exists before setting logo to None
            if context.get("client") is None:
                context["client"] = {}
            context["client"]["logo"] = None

        # Fix client diffusion_list if it's None
        if (
            context["client"].get("diffusion_list") is None
            or context["client"]["diffusion_list"] == ""
        ):
            context["client"]["diffusion_list"] = []

        # Additional safety checks for nested objects
        if context.get("assessment_structure"):
            if isinstance(context["assessment_structure"], dict):
                if context["assessment_structure"].get("s_fields") is None:
                    context["assessment_structure"]["s_fields"] = {}
                if context["assessment_structure"].get("a_fields") is None:
                    context["assessment_structure"]["a_fields"] = {}

        # Safety check for vulnerability fields
        for vuln in context.get("vulnerabilities", []):
            if isinstance(vuln, dict):
                if vuln.get("fields") is None:
                    vuln["fields"] = {}

        try:

            # Create InlineImage objects just before rendering to avoid doc corruption
            for chart_type in ["pie", "bar"]:
                chart_bytes_key = f"{chart_type}_chart_bytes"
                chart_size_key = f"{chart_type}_chart_size"
                if chart_bytes_key in context:
                    try:
                        context[f"{chart_type}_chart"] = InlineImage(
                            doc,
                            io.BytesIO(context[chart_bytes_key]),
                            width=Mm(context[chart_size_key]),
                        )
                        # Clean up the temporary bytes
                        del context[chart_bytes_key]
                        del context[chart_size_key]
                    except Exception as img_error:
                        context[f"{chart_type}_chart"] = None

            doc.render(context, jinja_env)

            # Add Internal Linking workaround
            update_internal_links(doc)

            # Save and return
            report_stream = io.BytesIO()
            doc.save(report_stream)

            report_stream.seek(0)
            response = HttpResponse(
                report_stream.getvalue(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = (
                'attachment; filename="generated_report.docx"'
            )
            return response

        except Exception as e:
            raise Exception(f"Error generating report: {str(e)}")

    except Exception as e:
        raise Exception(f"Error setting up report generation: {str(e)}")


# doc.save(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'generated_report.docx'))
# ------- Upadate table of contents usign windows api-------
# word = win32com.client.DispatchEx("Word.Application")
# try:
#     doc = word.Documents.Open(os.path.join(os.path.dirname(os.path.abspath(__file__)), 'generated_report.docx'))
#     doc.TablesOfContents(1).Update()
#     doc.Close(SaveChanges=True)
#     word.Quit()
#     del word
# except:
#     print("error")
#     del word


# ------- Remove white spaces inside all the tables -------
# Remove white spaces inside all the tables (that are not hyperlink)
# for table in doc.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             for paragraph in cell.paragraphs:
#                 if len(paragraph.text) == 0:
#                     if "hyperlink" not in str(paragraph._p.xml):
#                         p = paragraph._element
#                         p.getparent().remove(p)
#                         p._p = p._element = None


# ------- Rendering Summary, Scope, Notes (html -> doc) -------
# context["summary"] = doc.new_subdoc(html_2_sub_docx(context["summary"]))
# context["notes"] = doc.new_subdoc(html_2_sub_docx(context["notes"]))
# context["scope"] = doc.new_subdoc(html_2_sub_docx(context["scope"]))


# ------- Rendering vulnerability POC images -------
# for vuln in context['vulnerabilities']:
#     for poc_screenshot in vuln['poc_screenshots']:
#         poc_screenshot['image'] = InlineImage(doc, poc_screenshot['image'],width=Mm(140))


# ------- Rendering Attached Files -------
# images = ["jpg", "jpeg", "png"]
# for file in context['files']:
#     if file['name'].split(".")[-1] in images:
#         file['file'] = InlineImage(doc, file['file'],width=Mm(70))


# ------- Rendering old custom fields -------
# if context['fields']:
#     for field_name2, field_value2 in context['fields'].items():
#         if field_name2[:4] == 'link':
#             link_tuple2 = field_value2.split(',')
#             if len(link_tuple2) == 2:
#                 link_name2 = link_tuple2[0].strip()
#                 link_url2 = link_tuple2[1].strip()
#                 rt2 = RichText()
#                 rt2.add(str(link_name2),url_id=doc.build_url_id(str(link_url2)))
#                 context['fields'][field_name2] = rt2

# ------- Rendering vuln custom fields -------
# if vuln['custom_fields'] is not None:
#     for field_name, field_value in vuln['custom_fields'].items():
#         if field_name[:4] == 'link':
#             link_tuple = field_value.split(',')
#             if len(link_tuple) == 2:
#                 link_name = link_tuple[0].strip()
#                 link_url = link_tuple[1].strip()
#                 rt = RichText()
#                 rt.add(str(link_name),url_id=doc.build_url_id(str(link_url)))
#                 vuln['custom_fields'][field_name] = rt
