
from docx.enum.dml import MSO_THEME_COLOR
import docx.oxml
import docx.oxml.ns as ns
import re

def add_bookmark(run, bookmark_name):
    ''' Adds a word bookmark to a run '''
    tag = run._r
    start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
    start.set(docx.oxml.ns.qn('w:id'), '0')
    start.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(start)

    text = docx.oxml.OxmlElement('w:r')
    tag.append(text)

    end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
    end.set(docx.oxml.ns.qn('w:id'), '0')
    end.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(end)

    return run

def add_internal_hyperlink(paragraph, link_to, text):
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('w:anchor'), link_to, )
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    r = paragraph.add_run ()
    r._r.append (hyperlink)
    r.font.name = "Calibri"
    r.font.color.theme_color = MSO_THEME_COLOR.HYPERLINK
    r.font.underline = True    


def update_internal_links(doc):
    # Add bookmarks using their numbers as names
    for paragraph in doc.paragraphs:
        if  '++bookmark' in paragraph.text :
            bookmark_number = re.findall(r'\+\+bookmark(\d+)', paragraph.text)
            bookmark_number = str(bookmark_number[0])
            
            paragraph.text = re.sub(r'\+\+bookmark\d+', ' ', paragraph.text)
            add_bookmark(paragraph.add_run(text=None, style=None),bookmark_number)


    # Look for links to the bookmark's names in all the paragraphs
    for paragraph in doc.paragraphs:                 
        if  'link_to_bookmark' in paragraph.text :
            bookmark_number = re.findall(r'link_to_bookmark(\d+)', paragraph.text)
            bookmark_number = str(bookmark_number[0])
            
            updated_text = re.sub(r'link_to_bookmark\d+', ' ', paragraph.text)
            paragraph.text = ''
            add_internal_hyperlink(paragraph, bookmark_number, updated_text)


    # Look for links to the bookmark's names in all the tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if  'link_to_bookmark' in paragraph.text :
                        bookmark_number = re.findall(r'link_to_bookmark(\d+)', paragraph.text)
                        bookmark_number = str(bookmark_number[0])
                        
                        updated_text = re.sub(r'link_to_bookmark\d+', ' ', paragraph.text)
                        paragraph.text = ''
                        add_internal_hyperlink(paragraph, bookmark_number, updated_text)
                        
    return doc


